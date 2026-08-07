import io
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Optional, Tuple
import logging

logger = logging.getLogger(__name__)


class DocumentParserService:
    """
    Enterprise Multi-format Document Parser supporting PDF, DOCX, TXT, and Markdown.
    Includes page range filtering, encoding fallbacks, and smart adaptive chunking.
    """

    @classmethod
    def extract_text(
        cls,
        file_bytes: bytes,
        filename: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, int]:
        """
        Detects file extension and extracts structured text and total pages.
        Returns: (extracted_text, total_pages)
        """
        ext = filename.lower().split(".")[-1] if "." in filename else ""

        if ext == "pdf":
            return cls.extract_text_from_pdf(file_bytes, filename, start_page, end_page)
        elif ext in ["docx", "doc"]:
            return cls.extract_text_from_docx(file_bytes, filename)
        elif ext in ["txt", "md", "markdown", "csv"]:
            return cls.extract_text_from_plain_text(file_bytes, filename)
        else:
            # Fallback text decoding
            return cls.extract_text_from_plain_text(file_bytes, filename)

    @classmethod
    def extract_text_from_pdf(
        cls,
        file_bytes: bytes,
        filename: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, int]:
        """
        Extracts text from PDF with page numbering annotations and page range support.
        """
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            total_pages = len(reader.pages)

            if total_pages == 0:
                return "", 0

            # Normalize 1-indexed page bounds
            p_start = max(0, (start_page - 1)) if start_page and start_page > 0 else 0
            p_end = min(total_pages, end_page) if end_page and end_page > 0 else total_pages

            if p_start >= total_pages:
                p_start = 0
            if p_end <= p_start:
                p_end = total_pages

            extracted_chunks = []
            for idx in range(p_start, p_end):
                try:
                    page = reader.pages[idx]
                    page_text = page.extract_text() or ""
                    clean_text = page_text.strip()
                    if clean_text:
                        extracted_chunks.append(f"--- [Trang {idx + 1}/{total_pages}] ---\n{clean_text}")
                except Exception as page_err:
                    logger.warning(f"Lỗi đọc trang {idx + 1} của {filename}: {page_err}")

            full_text = "\n\n".join(extracted_chunks).strip()
            return full_text, total_pages
        except Exception as e:
            logger.error(f"Lỗi phân tích PDF {filename}: {e}")
            # Fallback to binary decode
            return cls.extract_text_from_plain_text(file_bytes, filename)

    @classmethod
    def extract_text_from_docx(cls, file_bytes: bytes, filename: str) -> Tuple[str, int]:
        """
        Extracts text from DOCX files without heavy external dependencies
        using python-docx if installed or native OpenXML extraction as fallback.
        """
        try:
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_content = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_content:
                            tables_text.append(row_content)
                combined = "\n".join(paragraphs + tables_text)
                return combined.strip(), max(1, len(combined) // 2500)
            except ImportError:
                # Native OpenXML zip parser fallback
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as docx_zip:
                    xml_content = docx_zip.read("word/document.xml")
                    tree = ET.fromstring(xml_content)
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    paragraphs = []
                    for p in tree.findall(".//w:p", ns):
                        texts = [node.text for node in p.findall(".//w:t", ns) if node.text]
                        if texts:
                            paragraphs.append("".join(texts))
                    combined = "\n".join(paragraphs).strip()
                    return combined, max(1, len(combined) // 2500)
        except Exception as e:
            logger.error(f"Lỗi phân tích DOCX {filename}: {e}")
            return cls.extract_text_from_plain_text(file_bytes, filename)

    @classmethod
    def extract_text_from_plain_text(cls, file_bytes: bytes, filename: str) -> Tuple[str, int]:
        """
        Decodes text files using multiple encoding strategies.
        """
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text = file_bytes.decode(encoding)
                clean_text = text.strip()
                pages_est = max(1, len(clean_text) // 2500)
                return clean_text, pages_est
            except UnicodeDecodeError:
                continue
        # Last resort ignore errors
        text = file_bytes.decode("utf-8", errors="ignore").strip()
        return text, max(1, len(text) // 2500)

    @classmethod
    def smart_chunk_text(cls, text: str, num_questions: int) -> str:
        """
        Adaptive smart chunking:
        - Allocates between 10,000 and 25,000 characters based on question volume.
        - Ensures fast LLM processing (<4s) while maintaining comprehensive pedagogical context.
        """
        if not text:
            return ""

        # Dynamic character budget: base 10k + 1k per question, capped at 25k
        char_limit = min(25000, max(10000, num_questions * 1200))

        if len(text) <= char_limit:
            return text

        # If text exceeds budget, preserve beginning, middle, and end segments cleanly
        header_budget = int(char_limit * 0.4)
        middle_budget = int(char_limit * 0.4)
        footer_budget = char_limit - header_budget - middle_budget

        mid_point = len(text) // 2
        mid_start = max(0, mid_point - (middle_budget // 2))

        chunk = (
            text[:header_budget]
            + "\n\n... [NỘI DUNG GIỮA TÀI LIỆU] ...\n\n"
            + text[mid_start : mid_start + middle_budget]
            + "\n\n... [NỘI DUNG CUỐI TÀI LIỆU] ...\n\n"
            + text[-footer_budget:]
        )
        return chunk
